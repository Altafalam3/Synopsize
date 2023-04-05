import docx
from docx.oxml.ns import qn
from tkinter import filedialog

def wordDoc(dayOfMonth, time, dayOfWeek, month, title, summary):

    # Define the variables
    day_of_month = str(dayOfMonth)
    time = str(time)
    day_of_week = str(dayOfWeek)
    month = str(month)
    title = str(title)
    summary = str(summary)

    # "This is a summary of my document. It contains important information that I want to convey to the reader. Sarah is a sentence web developer. Altaf is a math genius. Amisha is Ausha. Abhigyan only builds cool stuff."

    # Create a new Word document
    document = docx.Document()

    # Add the date and time to the header
    header = document.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_paragraph.add_run()
    header_run.add_text(f"{day_of_month} {month}").bold = True
    header_run.add_text("\t\t\t\t\t\t\t")
    header_run.add_text(f"{day_of_week}, {time}").bold = True

    # Add the title to the document
    title_paragraph = document.add_paragraph(title)
    title_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.style = "Title"


    # Add the summary to the document
    summary_paragraph = document.add_paragraph()
    sentences = summary.split(".")
    for sentence in sentences:
        if sentence:
            summary_bullet_points = summary_paragraph.add_run(f"\n• {sentence.strip()}.")
            summary_bullet_points.font.name = "Symbol"
            summary_bullet_points._element.rPr.rFonts.set(qn('w:eastAsia'), 'Symbol')
            summary_bullet_points._element.rPr.rFonts.set(qn('w:cs'), 'Symbol')


    # Ask the user for the file path to save the document
    filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=(("Word Document", "*.docx"), ("All files", "*.*")))

    # Save the document to the specified file path
    document.save(filename)
